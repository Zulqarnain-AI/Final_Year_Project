#!/usr/bin/env python3
"""
Generate a comprehensive models performance report in .docx format
Covers both Audio Model and Symptoms Model
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============= TITLE PAGE =============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Final Year Project\n')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 102, 204)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Respiratory Disease Diagnosis System\n')
subtitle_run.font.size = Pt(18)
subtitle_run.font.bold = True

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2_run = subtitle2.add_run('Integrated Models Performance Report\n\n')
subtitle2_run.font.size = Pt(14)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
date_run.font.size = Pt(11)
date_run.italic = True

doc.add_paragraph()  # Spacing

# ============= TABLE OF CONTENTS =============
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Architecture Overview',
    '3. Audio Model Performance (CNN)',
    '4. Symptoms Model Performance (Random Forest)',
    '5. Integrated Diagnosis Pipeline',
    '6. Conclusion and Recommendations'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============= SECTION 1: EXECUTIVE SUMMARY =============
doc.add_heading('1. Executive Summary', level=1)

summary_text = """
This report provides a comprehensive analysis of the two machine learning models that form the core of the Respiratory Disease Diagnosis System. The system integrates:

• An Audio Model that analyzes breathing patterns from recorded audio samples
• A Symptoms Model that processes patient-reported symptoms along with demographic data

Both models work in conjunction to provide accurate respiratory disease diagnosis with high confidence scores. The system can diagnose five respiratory conditions: Asthma, COPD, Bronchitis, Pneumonia, and Healthy.
"""
doc.add_paragraph(summary_text)

# ============= SECTION 2: SYSTEM ARCHITECTURE =============
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 System Components', level=2)

architecture_text = """
The respiratory disease diagnosis system consists of two independent classification pipelines that operate in parallel:

Input Layer:
  • Audio Input: 5-second audio samples (WAV, MP3, FLAC, OGG, M4A formats)
  • Patient Data: Symptoms (text), Age (numeric), Sex (categorical)

Processing Layer:
  • Audio Preprocessing: MFCC extraction using librosa (22,050 Hz sample rate, 128 mel-spectrogram bins)
  • Data Preprocessing: TF-IDF vectorization for symptoms, standardization for age

Model Layer:
  • Audio Model: Convolutional Neural Network (CNN) with deep learning
  • Symptoms Model: Random Forest Classifier (200 estimators)

Output Layer:
  • Predictions: Disease class label with confidence score
  • Integration: Both models' predictions are combined for final diagnosis
"""
doc.add_paragraph(architecture_text)

# ============= SECTION 3: AUDIO MODEL =============
doc.add_heading('3. Audio Model Performance (CNN)', level=1)

doc.add_heading('3.1 Model Architecture', level=2)

audio_arch_text = """
Model Type: Convolutional Neural Network (CNN)
Framework: TensorFlow/Keras
Input Shape: (128, 430, 1) - Mel-spectrogram features
Number of Classes: 5 (asthma, copd, bronchial, pneumonia, healthy)

Architecture Layers:
  • Input Layer: Accepts mel-spectrogram features
  • Convolutional Blocks: Multiple Conv2D layers with ReLU activation
  • Pooling Layers: MaxPooling2D for feature dimensionality reduction
  • Batch Normalization: Stabilizes training and improves generalization
  • Dropout Layers: Prevents overfitting with 50% dropout rate
  • Flatten Layer: Converts 2D feature maps to 1D vector
  • Dense Layers: Fully connected layers with 256 and 128 units
  • Output Layer: Softmax activation for multi-class classification
"""
doc.add_paragraph(audio_arch_text)

doc.add_heading('3.2 Training Configuration', level=2)

audio_train_table = doc.add_table(rows=11, cols=2)
audio_train_table.style = 'Light Grid Accent 1'
header_cells = audio_train_table.rows[0].cells
header_cells[0].text = 'Parameter'
header_cells[1].text = 'Value'

training_params = [
    ('Optimizer', 'Adam with learning rate 1e-4'),
    ('Loss Function', 'Categorical Crossentropy (label smoothing: 0.1)'),
    ('Batch Size', '16'),
    ('Epochs', '50'),
    ('Train/Validation Split', '80/20'),
    ('Sample Rate', '22,050 Hz'),
    ('Duration', '5 seconds'),
    ('Mel-Spectrogram Bins', '128'),
    ('Callbacks', 'EarlyStopping (patience=15), ReduceLROnPlateau'),
    ('Data Augmentation', 'Random noise injection, pitch shifting')
]

for i, (param, value) in enumerate(training_params, 1):
    row = audio_train_table.rows[i].cells
    row[0].text = param
    row[1].text = value

doc.add_heading('3.3 Performance Metrics', level=2)

audio_perf_text = """
The audio model evaluates respiratory disease classification using:

Metrics Evaluated:
  • Accuracy: Overall percentage of correct predictions
  • Precision: True positive rate among predicted positives
  • Recall: True positive rate among actual positives
  • F1-Score: Harmonic mean of precision and recall
  • Confusion Matrix: Detailed breakdown of prediction errors

Evaluation Process:
  • Validation Set: 20% of balanced dataset used for evaluation
  • Batch Prediction: Model.predict() on validation set
  • Classification Report: Per-class metrics for all 5 disease categories
"""
doc.add_paragraph(audio_perf_text)

doc.add_heading('3.4 Strengths and Weaknesses', level=2)

strengths = doc.add_paragraph('Strengths:', style='Heading 3')
doc.add_paragraph('Deep learning captures complex patterns in audio spectrograms', style='List Bullet')
doc.add_paragraph('Handles frequency domain features effectively for respiratory analysis', style='List Bullet')
doc.add_paragraph('Batch normalization and dropout prevent overfitting', style='List Bullet')
doc.add_paragraph('Early stopping prevents unnecessary training iterations', style='List Bullet')

weaknesses = doc.add_paragraph('Weaknesses:', style='Heading 3')
doc.add_paragraph('Requires large amounts of quality audio data for training', style='List Bullet')
doc.add_paragraph('Sensitive to audio quality, background noise, and recording device', style='List Bullet')
doc.add_paragraph('Longer inference time compared to traditional ML models', style='List Bullet')
doc.add_paragraph('Black-box nature makes interpretation difficult', style='List Bullet')

# ============= SECTION 4: SYMPTOMS MODEL =============
doc.add_heading('4. Symptoms Model Performance (Random Forest)', level=1)

doc.add_heading('4.1 Model Architecture', level=2)

symptoms_arch_text = """
Model Type: Random Forest Classifier
Framework: scikit-learn
Number of Estimators: 200 decision trees
Input Features: Symptoms (text), Age (numeric), Sex (categorical)
Number of Classes: 5 respiratory diseases

Feature Engineering:
  • Symptoms: TF-IDF Vectorization (term frequency-inverse document frequency)
  • Age: StandardScaler normalization
  • Sex: OneHotEncoder categorical encoding
  • Pipeline: Scikit-learn Pipeline for reproducible preprocessing
"""
doc.add_paragraph(symptoms_arch_text)

doc.add_heading('4.2 Model Configuration', level=2)

symptoms_train_table = doc.add_table(rows=11, cols=2)
symptoms_train_table.style = 'Light Grid Accent 1'
header_cells = symptoms_train_table.rows[0].cells
header_cells[0].text = 'Parameter'
header_cells[1].text = 'Value'

symptoms_params = [
    ('Algorithm', 'Random Forest Classifier'),
    ('Number of Trees', '200'),
    ('Max Depth', 'None (unlimited)'),
    ('Min Samples Split', '2'),
    ('Class Weighting', 'Balanced (handles imbalanced data)'),
    ('Train/Test Split', '80/20'),
    ('Cross-Validation', '5-Fold Stratified K-Fold'),
    ('Feature Scaling', 'StandardScaler for numeric features'),
    ('Categorical Encoding', 'OneHotEncoder for categorical features'),
    ('Text Vectorization', 'TF-IDF (Term Frequency-Inverse Document Frequency)')
]

for i, (param, value) in enumerate(symptoms_params, 1):
    row = symptoms_train_table.rows[i].cells
    row[0].text = param
    row[1].text = value

doc.add_heading('4.3 Cross-Validation Results', level=2)

cv_text = """
The model was extensively validated using 5-Fold Stratified Cross-Validation:

Evaluation Metrics:
  • Accuracy: Proportion of correct predictions across all folds
  • Precision (Weighted): Accuracy among predicted disease cases
  • Recall (Weighted): Ability to find actual disease cases
  • F1 Score (Weighted): Balance between precision and recall

Class Imbalance Handling:
  • Balanced class weights are automatically computed
  • Prevents bias toward majority disease class
  • Improves minority class prediction performance
  • Stratified K-Fold ensures class distribution in all folds
"""
doc.add_paragraph(cv_text)

doc.add_heading('4.4 Model Strengths and Limitations', level=2)

strengths_sym = doc.add_paragraph('Strengths:', style='Heading 3')
doc.add_paragraph('Interpretable: Feature importance scores show which symptoms matter most', style='List Bullet')
doc.add_paragraph('Robust to outliers in patient age and demographics', style='List Bullet')
doc.add_paragraph('Handles imbalanced datasets with balanced class weights', style='List Bullet')
doc.add_paragraph('Fast training and inference suitable for real-time diagnosis', style='List Bullet')
doc.add_paragraph('Doesn\'t require normalization of input data', style='List Bullet')

weaknesses_sym = doc.add_paragraph('Limitations:', style='Heading 3')
doc.add_paragraph('Depends on accurate patient symptom reporting', style='List Bullet')
doc.add_paragraph('Limited by vocabulary in symptom dataset', style='List Bullet')
doc.add_paragraph('Age information may not capture all relevant patient context', style='List Bullet')
doc.add_paragraph('Cannot detect new symptom patterns not seen in training data', style='List Bullet')

# ============= SECTION 5: INTEGRATED PIPELINE =============
doc.add_heading('5. Integrated Diagnosis Pipeline', level=1)

doc.add_heading('5.1 Diagnosis Workflow', level=2)

workflow_text = """
The final diagnosis system integrates both models into a unified pipeline:

Step 1: Input Collection
  • Patient selects symptoms from predefined list
  • Patient records or uploads 5-second breathing audio
  • System captures patient age and sex

Step 2: Parallel Processing
  • Audio Model: Processes the audio sample through CNN
    - Outputs: Disease prediction with confidence score
  • Symptoms Model: Processes patient data through Random Forest
    - Outputs: Disease prediction with confidence score

Step 3: Prediction Fusion
  • Combine predictions from both models
  • Weight predictions based on model confidence scores
  • Generate ensemble prediction

Step 4: Output Generation
  • Display final diagnosis with confidence percentage
  • Provide detailed analysis from both models
  • Generate diagnostic report with timestamps
  • Store results in patient medical records

Validation Requirements:
  • Both audio and symptoms must be provided
  • Audio must be in supported format (WAV, MP3, FLAC, OGG, M4A)
  • Patient must select at least one symptom
  • Authentication required (user must be logged in)
"""
doc.add_paragraph(workflow_text)

doc.add_heading('5.2 Error Handling and Validation', level=2)

validation_table = doc.add_table(rows=6, cols=3)
validation_table.style = 'Light Grid Accent 1'
header_cells = validation_table.rows[0].cells
header_cells[0].text = 'Validation Check'
header_cells[1].text = 'Condition'
header_cells[2].text = 'Action'

validation_checks = [
    ('Audio File', 'Missing or unsupported format', 'Display error, request valid audio'),
    ('Symptoms', 'None selected', 'Display error, require symptom selection'),
    ('Authentication', 'User not logged in', 'Redirect to login page'),
    ('Audio Quality', 'Too short or corrupted', 'Show warning, allow retry'),
    ('Processing Timeout', 'Analysis takes too long', 'Notify user, suggest retry')
]

for i, (check, condition, action) in enumerate(validation_checks, 1):
    row = validation_table.rows[i].cells
    row[0].text = check
    row[1].text = condition
    row[2].text = action

# ============= SECTION 6: CONCLUSION =============
doc.add_heading('6. Conclusion and Recommendations', level=1)

doc.add_heading('6.1 System Performance Summary', level=2)

summary_conclusion = """
The integrated respiratory disease diagnosis system demonstrates a multi-modal approach to medical diagnosis:

Audio Model (CNN):
  • Specialized in analyzing acoustic patterns in breathing
  • Captures temporal and frequency domain features
  • Achieves robust performance across 5 disease categories
  • Requires high-quality audio input for optimal performance

Symptoms Model (Random Forest):
  • Leverages structured patient medical data
  • Fast, interpretable, and production-ready
  • Handles class imbalance effectively
  • Improves diagnosis accuracy through ensemble approach

Integrated System:
  • Combines complementary strengths of both models
  • Provides dual validation for increased confidence
  • Reduces dependency on single data modality
  • Improves overall diagnostic accuracy
"""
doc.add_paragraph(summary_conclusion)

doc.add_heading('6.2 Recommendations for Improvement', level=2)

recommendations = [
    ('Data Collection', 'Expand audio dataset with diverse patient populations and recording devices'),
    ('Model Enhancement', 'Experiment with ensemble methods combining CNN and RandomForest directly'),
    ('Feature Engineering', 'Incorporate additional medical features (vital signs, medical history)'),
    ('Real-time Monitoring', 'Track model performance in production and retrain periodically'),
    ('User Testing', 'Conduct clinical validation studies with healthcare professionals'),
    ('Privacy', 'Implement encryption for storing audio and patient medical data'),
    ('Explainability', 'Add SHAP values and LIME for model prediction interpretation'),
    ('Reliability', 'Implement confidence thresholds for uncertain predictions')
]

for i, (category, recommendation) in enumerate(recommendations, 1):
    para = doc.add_paragraph(style='List Number')
    para.add_run(f'{category}: ').bold = True
    para.add_run(recommendation)

doc.add_heading('6.3 Deployment Considerations', level=2)

deployment_text = """
For production deployment of this system:

Infrastructure:
  • Backend API: Flask/FastAPI for model serving
  • Audio Processing: Real-time audio preprocessing pipeline
  • Database: Secure storage for patient records and predictions
  • Frontend: React-based user interface for patient interaction

Compliance:
  • HIPAA/GDPR compliance for medical data handling
  • Model validation against clinical standards
  • Regular audits and performance monitoring
  • Clear disclaimers about AI limitations

Performance Metrics to Monitor:
  • Model accuracy on new patient data
  • Inference latency (< 2 seconds for real-time diagnosis)
  • False positive and false negative rates
  • User feedback and clinical validation results
"""
doc.add_paragraph(deployment_text)

doc.add_page_break()

# ============= APPENDIX =============
doc.add_heading('Appendix: Technical Specifications', level=1)

doc.add_heading('A. Audio Processing Pipeline', level=2)

audio_pipeline = """
Feature Extraction Process:
  1. Load audio file at 22,050 Hz sample rate
  2. Pad or truncate to 110,250 samples (5 seconds)
  3. Apply data augmentation:
     • Random noise injection (±0.005 amplitude)
     • Pitch shifting (±2 semitones)
     • Time stretching (±5% speed)
  4. Extract mel-spectrogram:
     • 128 mel-frequency bins
     • Window size: 2048 samples
     • Hop length: 512 samples
  5. Normalize to [0, 1] range
  6. Reshape to (128, 430, 1) for CNN input
"""
doc.add_paragraph(audio_pipeline)

doc.add_heading('B. Symptoms Processing Pipeline', level=2)

symptoms_pipeline = """
Feature Preprocessing Steps:
  1. Data Cleaning:
     • Remove rows with missing symptoms
     • Fill age with median value
     • Fill sex with mode value
  2. Text Preprocessing:
     • Convert symptoms to lowercase
     • Remove extra whitespace
     • Vectorize using TF-IDF
  3. Categorical Encoding:
     • OneHotEncode sex (Male/Female)
  4. Numeric Scaling:
     • StandardScale age to zero mean, unit variance
  5. Combine all features for Random Forest input
"""
doc.add_paragraph(symptoms_pipeline)

doc.add_heading('C. Class Definitions', level=2)

diseases_table = doc.add_table(rows=6, cols=2)
diseases_table.style = 'Light Grid Accent 1'
header_cells = diseases_table.rows[0].cells
header_cells[0].text = 'Disease'
header_cells[1].text = 'Description'

diseases = [
    ('Asthma', 'Chronic inflammatory disorder of airways, causes wheezing and shortness of breath'),
    ('COPD', 'Chronic Obstructive Pulmonary Disease, progressive lung disease with limited airflow'),
    ('Bronchitis', 'Inflammation of bronchial tubes, causes persistent cough and mucus production'),
    ('Pneumonia', 'Infection causing lung inflammation, consolidation, and fluid in alveoli'),
    ('Healthy', 'Normal respiratory function without signs of disease')
]

for i, (disease, desc) in enumerate(diseases, 1):
    row = diseases_table.rows[i].cells
    row[0].text = disease
    row[1].text = desc

# ============= FOOTER =============
doc.add_page_break()
doc.add_heading('Document Information', level=2)

footer_info = f"""
Report Title: Respiratory Disease Diagnosis System - Models Performance Report
Generated Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Project: Final Year Project - Integrated Audio and Symptom-Based Diagnosis
System Components:
  • Audio Model: CNN-based deep learning classifier
  • Symptoms Model: Random Forest ensemble classifier
  • Frontend: React-based patient interface
  • Backend: Flask API with TensorFlow/scikit-learn models

For questions or additional information, please refer to the project documentation
and source code repository.
"""
doc.add_paragraph(footer_info)

# ============= SAVE DOCUMENT =============
output_path = 'Models_Performance_Report.docx'
doc.save(output_path)
print(f"✅ Document generated successfully: {output_path}")
print(f"📄 File size: {__import__('os').path.getsize(output_path) / 1024:.2f} KB")
